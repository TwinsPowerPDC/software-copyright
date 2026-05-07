#!/usr/bin/env python3
"""源代码文档生成脚本 — 读取文件列表，生成 A4 格式 Word 文档。"""

import argparse
import os
import sys
from collections import Counter
from pathlib import Path


def parse_args():
    parser = argparse.ArgumentParser(
        description='将源代码文件汇编为软著申请用的 Word 文档'
    )
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument(
        '--files-from', type=Path, metavar='FILE',
        help='后端文件列表路径（每行一个绝对路径），全量纳入'
    )
    group.add_argument(
        '--files', nargs='+', type=Path, metavar='PATH',
        help='直接传入后端源代码文件路径'
    )
    parser.add_argument(
        '--frontend-files-from', type=Path, metavar='FILE',
        help='前端文件列表路径（每行一个绝对路径），按行数降序加入直到满足双条件'
    )
    parser.add_argument(
        '--frontend-min-lines', type=int, default=1000, metavar='N',
        help='前端最低保障行数，默认 1000'
    )
    parser.add_argument(
        '--all-files-from', type=Path, metavar='FILE',
        help='项目所有源代码文件路径清单（每行一个），用于统计代码总行数'
    )
    parser.add_argument(
        '--output', required=True, type=Path, metavar='DOCX',
        help='输出 .docx 文件路径'
    )
    parser.add_argument(
        '--min-lines', type=int, default=3500, metavar='N',
        help='总行数最低目标，默认 3500'
    )
    return parser.parse_args()


def load_file_list(args):
    """从参数解析文件列表，返回路径字符串列表。"""
    if args.files_from:
        if not args.files_from.exists():
            print(f'错误：文件列表不存在：{args.files_from}', file=sys.stderr)
            sys.exit(1)
        paths = [
            line.strip() for line in args.files_from.read_text(encoding='utf-8').splitlines()
            if line.strip()
        ]
    else:
        paths = [str(p) for p in args.files]

    if not paths:
        print('错误：文件列表为空', file=sys.stderr)
        sys.exit(1)
    return paths


def count_total_lines(file_list_path):
    """遍历文件路径清单，逐文件统计行数求和，返回总行数。"""
    total = 0
    for line in file_list_path.read_text(encoding='utf-8').splitlines():
        fp = line.strip()
        if not fp:
            continue
        try:
            with open(fp, 'r', encoding='utf-8', errors='replace') as f:
                total += sum(1 for _ in f)
        except Exception:
            pass
    return total


def compute_labels(file_paths):
    """为文件列表计算最简标识：唯一文件名直接用，重名时逐级补父目录直到区分开。

    示例：
      .../backend/SalesClueService.java  →  SalesClueService.java  （唯一）
      .../components/details/index.tsx   →  details/index.tsx       （重名，补一级父目录）
      .../components/list/index.tsx      →  list/index.tsx          （重名，补一级父目录）
    """
    basenames = [os.path.basename(fp) for fp in file_paths]
    name_counts = Counter(basenames)
    labels = {}
    for fp in file_paths:
        basename = os.path.basename(fp)
        if name_counts[basename] == 1:
            labels[fp] = basename
        else:
            parts = Path(fp).parts
            for depth in range(2, len(parts) + 1):
                candidate = os.path.join(*parts[-depth:])
                unique = all(
                    os.path.join(*Path(other).parts[-depth:]) != candidate
                    for other in file_paths
                    if other != fp and os.path.basename(other) == basename
                )
                if unique:
                    labels[fp] = candidate
                    break
            else:
                labels[fp] = basename
    return labels


def set_run_font(run, font_name='Microsoft YaHei Light', font_size_pt=10, color_hex='000000'):
    """设置 run 的字体（Latin + 东亚 + CS）、字号和颜色。"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_doc_defaults(doc, font_name='Microsoft YaHei Light', font_size_pt=10, color_hex='000000'):
    """设置文档级默认字体（w:docDefaults > w:rPrDefault）和 Normal 样式字体。"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 设置 Normal 样式
    normal_style = doc.styles['Normal']
    normal_style.font.name = font_name
    normal_style.font.size = Pt(font_size_pt)
    normal_style.font.color.rgb = RGBColor.from_string(color_hex)
    rPr = normal_style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

    # 设置文档级 docDefaults
    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is None:
        doc_defaults = OxmlElement('w:docDefaults')
        styles_elem.insert(0, doc_defaults)

    rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        doc_defaults.insert(0, rPrDefault)

    rPr_def = rPrDefault.find(qn('w:rPr'))
    if rPr_def is None:
        rPr_def = OxmlElement('w:rPr')
        rPrDefault.append(rPr_def)

    rFonts_def = rPr_def.find(qn('w:rFonts'))
    if rFonts_def is None:
        rFonts_def = OxmlElement('w:rFonts')
        rPr_def.insert(0, rFonts_def)
    rFonts_def.set(qn('w:ascii'), font_name)
    rFonts_def.set(qn('w:hAnsi'), font_name)
    rFonts_def.set(qn('w:eastAsia'), font_name)
    rFonts_def.set(qn('w:cs'), font_name)

    color_elem = rPr_def.find(qn('w:color'))
    if color_elem is None:
        color_elem = OxmlElement('w:color')
        rPr_def.append(color_elem)
    color_elem.set(qn('w:val'), color_hex)


def read_lines(file_path):
    """读取文件内容，返回行列表；失败返回空列表。"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.readlines()
    except Exception as e:
        print(f'读取失败: {file_path}: {e}', file=sys.stderr)
        return []


def build_document(backend_paths, frontend_paths, output_path,
                   min_lines=3500, frontend_min_lines=1000, total_lines=None):
    """生成 Word 文档。

    Args:
        backend_paths: 后端文件路径列表（按优先级顺序，全量纳入）
        frontend_paths: 前端文件路径列表（脚本内部按行数降序排列）
        output_path: 输出 .docx 路径
        min_lines: 总行数最低目标（默认 3500）
        frontend_min_lines: 前端最低保障行数（默认 1000）
        total_lines: 项目源代码总行数（由 --all-files-from 统计），None 则不写备注
    """
    import docx
    from docx.shared import Pt, Inches

    doc = docx.Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)

    all_paths = backend_paths + frontend_paths
    labels = compute_labels(all_paths)

    # === 阶段一：后端全量纳入 ===
    included_files = []
    backend_lines = 0
    for file_path in backend_paths:
        lines = read_lines(file_path)
        if not lines:
            continue
        backend_lines += len(lines)
        included_files.append((file_path, lines))
    print(f'后端全量纳入：{backend_lines} 行，共 {len(included_files)} 个文件', file=sys.stderr)

    # === 阶段二：前端按行数降序加入，双条件满足后停止 ===
    # 条件1：总行数 ≥ min_lines；条件2：前端贡献 ≥ frontend_min_lines
    frontend_lines = 0
    if frontend_paths:
        # 预读各前端文件行数，按降序排列
        frontend_with_size = []
        for fp in frontend_paths:
            lines = read_lines(fp)
            if lines:
                frontend_with_size.append((fp, lines))
        frontend_with_size.sort(key=lambda x: len(x[1]), reverse=True)

        for file_path, lines in frontend_with_size:
            total = backend_lines + frontend_lines
            if total >= min_lines and frontend_lines >= frontend_min_lines:
                print(f'双条件满足（总行数 {total} ≥ {min_lines}，前端 {frontend_lines} ≥ {frontend_min_lines}），停止', file=sys.stderr)
                break
            frontend_lines += len(lines)
            included_files.append((file_path, lines))

    doc_lines = backend_lines + frontend_lines
    print(f'前端纳入：{frontend_lines} 行；总计：{doc_lines} 行', file=sys.stderr)

    # === 写入备注行 ===
    if total_lines is not None:
        if doc_lines < total_lines:
            note_text = f'本系统代码总行数：{total_lines} 行；本文档截取代码行数：{doc_lines} 行。'
        else:
            note_text = f'本系统代码总行数：{total_lines} 行；已全部放入本文档。'

        p = doc.add_paragraph()
        run = p.add_run(note_text)
        set_run_font(run, font_size_pt=10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # 备注后空一行
        p = doc.add_paragraph()
        run = p.add_run('')
        set_run_font(run, font_size_pt=10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # === 写入代码 ===
    for file_path, lines in included_files:
        # 文件标识行
        label = labels.get(file_path, os.path.basename(file_path))
        p = doc.add_paragraph()
        run = p.add_run(label)
        set_run_font(run, font_size_pt=10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        # 代码内容（每行独立段落）
        for line in lines:
            p = doc.add_paragraph()
            run = p.add_run(line.rstrip('\n'))
            set_run_font(run, font_size_pt=9)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(11)

        # 文件间空行（添加空 run 以确保字体一致）
        p = doc.add_paragraph()
        run = p.add_run('')
        set_run_font(run, font_size_pt=9)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    print(f'共写入 {doc_lines} 行，约 {doc_lines // 42} 页', file=sys.stderr)
    doc.save(str(output_path))
    # 结果路径输出到 stdout
    print(str(output_path))


def main():
    args = parse_args()

    # 早期验证：输出目录存在
    output_dir = args.output.parent
    if not output_dir.exists():
        print(f'错误：输出目录不存在：{output_dir}', file=sys.stderr)
        sys.exit(1)

    backend_paths = load_file_list(args)

    # 加载前端文件列表（可选）
    frontend_paths = []
    if args.frontend_files_from:
        if not args.frontend_files_from.exists():
            print(f'错误：前端文件列表不存在：{args.frontend_files_from}', file=sys.stderr)
            sys.exit(1)
        frontend_paths = [
            line.strip()
            for line in args.frontend_files_from.read_text(encoding='utf-8').splitlines()
            if line.strip()
        ]

    # 统计项目源代码总行数
    total_lines = None
    if args.all_files_from:
        if not args.all_files_from.exists():
            print(f'错误：全量文件列表不存在：{args.all_files_from}', file=sys.stderr)
            sys.exit(1)
        total_lines = count_total_lines(args.all_files_from)
        print(f'项目源代码总行数：{total_lines}', file=sys.stderr)

    build_document(
        backend_paths, frontend_paths, args.output,
        min_lines=args.min_lines,
        frontend_min_lines=args.frontend_min_lines,
        total_lines=total_lines,
    )


if __name__ == '__main__':
    main()
