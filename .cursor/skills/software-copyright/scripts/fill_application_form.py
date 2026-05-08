#!/usr/bin/env python3
"""申请表填充脚本 — 读取 .docx 模板，替换 {变量} 占位符（共 6 处）。"""

import argparse
import sys
from pathlib import Path

EXPECTED_PLACEHOLDERS = 6


def parse_args():
    parser = argparse.ArgumentParser(
        description='填充软著申请表 Word 模板中的 {变量} 占位符'
    )
    parser.add_argument(
        '--template', required=True, type=Path, metavar='DOCX',
        help='输入 .docx 模板文件路径'
    )
    parser.add_argument(
        '--name', required=True,
        help='软件全称（含品牌名，不含版本号，如"影能科技智慧园区管理系统"）'
    )
    parser.add_argument(
        '--short-name', required=True,
        help='软件简称（3-10字）'
    )
    parser.add_argument(
        '--features', required=True,
        help='软件主要功能（100-300字）'
    )
    parser.add_argument(
        '--tech', required=True,
        help='软件技术特点（50-150字）'
    )
    parser.add_argument(
        '--dev-complete-date', required=True,
        help='开发完成日期（与申请表格式一致，如 2024年9月25日）'
    )
    parser.add_argument(
        '--first-publish-date', required=True,
        help='首次发表日期（与申请表格式一致；未发表等口径由用户确认后原样填入）'
    )
    parser.add_argument(
        '--output', required=True, type=Path, metavar='DOCX',
        help='输出 .docx 文件路径'
    )
    return parser.parse_args()


def replace_placeholder_in_paragraph(para, new_text):
    """段落级合并替换，正确处理跨 runs 的 {变量} 占位符。"""
    full_text = ''.join(run.text for run in para.runs)
    if '{变量}' not in full_text:
        return False
    replaced = full_text.replace('{变量}', new_text, 1)
    if para.runs:
        para.runs[0].text = replaced
        for run in para.runs[1:]:
            run.text = ''
    return True


def process_paragraphs(paragraphs, context_map):
    """遍历段落，按上下文关键词匹配并替换 {变量}。"""
    replaced_count = 0
    for para in paragraphs:
        full_text = ''.join(run.text for run in para.runs)
        if '{变量}' not in full_text:
            continue
        matched = False
        for keyword, field_name, value in context_map:
            if keyword in full_text:
                if replace_placeholder_in_paragraph(para, value):
                    print(f'  已替换：{field_name}', file=sys.stderr)
                    replaced_count += 1
                    matched = True
                    break
        if not matched:
            print(f'  未识别的 {{变量}} 段落，请手动核查：{full_text[:60]}', file=sys.stderr)
    return replaced_count


def xml_level_replace(template_path, field_values, output_path):
    """Fallback：直接操作 docx XML 替换跨 run 拆分的 {变量} 占位符。

    .doc → .docx 转换后，{变量} 常被拆成 3 个独立 run：
        <w:t>{</w:t>  <w:t>变量</w:t>  <w:t>}</w:t>
    python-docx 的段落遍历可能因嵌套表格而遗漏。此函数绕过 python-docx，
    直接在 XML 层面定位并替换。
    """
    import zipfile
    import re

    with zipfile.ZipFile(template_path, 'r') as z:
        content = z.read('word/document.xml').decode('utf-8')
        other_files = {
            name: z.read(name)
            for name in z.namelist()
            if name != 'word/document.xml'
        }

    # 定位所有 <w:t>变量</w:t>
    var_pattern = r'<w:t>变量</w:t>'
    var_matches = list(re.finditer(var_pattern, content))
    if not var_matches:
        print('XML 中也未找到 "变量" 文本，模板可能已损坏', file=sys.stderr)
        return 0

    # 按文档出现顺序，根据上下文关键词映射到对应替换值（与 patch 后模板一致）
    context_keywords = [
        (['软件名称', '全称'], '软件全称', field_values['软件全称']),
        (['软件简称'], '软件简称', field_values['软件简称']),
        (['开发完成日期'], '开发完成日期', field_values['开发完成日期']),
        (['首次发表日期'], '首次发表日期', field_values['首次发表日期']),
        (['主要功能'], '主要功能', field_values['主要功能']),
        (['技术特点'], '技术特点', field_values['技术特点']),
    ]

    replacements = []
    for vm in var_matches:
        # 往回找最近的 <w:t>{</w:t>
        before = content[:vm.start()]
        brace_open_pos = before.rfind('<w:t>{</w:t>')
        # 往后找最近的 <w:t>}</w:t>
        after_match = re.search(r'<w:t>\}</w:t>', content[vm.end():])
        brace_close_end = vm.end() + after_match.end() if after_match else -1

        # 提取上下文纯文本用于关键词匹配
        ctx_start = max(0, brace_open_pos - 2000)
        ctx_text = re.sub(r'<[^>]+>', '', content[ctx_start:brace_open_pos])

        label, value = '未识别', '{变量}'
        for keywords, name, val in context_keywords:
            if any(kw in ctx_text for kw in keywords):
                label, value = name, val
                break

        replacements.append((brace_open_pos, vm.start(), vm.end(), brace_close_end, label, value))

    # 从后往前替换，避免位置偏移
    replaced_count = 0
    for brace_open_pos, var_start, var_end, brace_close_end, label, value in reversed(replacements):
        if brace_open_pos < 0 or brace_close_end < 0:
            print(f'  跳过：{label}（未找到配对大括号）', file=sys.stderr)
            continue
        brace_close_start = brace_close_end - len('<w:t>}</w:t>')
        # 清空 }
        content = content[:brace_close_start] + '<w:t></w:t>' + content[brace_close_end:]
        # 清空 变量
        content = content[:var_start] + '<w:t></w:t>' + content[var_end:]
        # 将 { 替换为实际值
        old_open = '<w:t>{</w:t>'
        new_open = f'<w:t xml:space="preserve">{value}</w:t>'
        content = content[:brace_open_pos] + new_open + content[brace_open_pos + len(old_open):]
        print(f'  已替换（XML）：{label}', file=sys.stderr)
        replaced_count += 1

    # 写回 docx
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        zout.writestr('word/document.xml', content.encode('utf-8'))
        for name, data in other_files.items():
            zout.writestr(name, data)

    return replaced_count



def fix_libreoffice_borders(docx_path):
    """修正 LibreOffice .doc→.docx 转换导致的边框样式错误。

    LibreOffice 会将 .doc 中的标准细线边框错误转换为 thickThinLargeGap + C0C0C0，
    本函数将其修正回 single + auto（标准黑色细线）。
    """
    import zipfile
    import re

    with zipfile.ZipFile(docx_path, 'r') as z:
        content = z.read('word/document.xml').decode('utf-8')
        entries = [(info, z.read(info.filename)) for info in z.infolist()
                   if info.filename != 'word/document.xml']
        doc_info = next(i for i in z.infolist() if i.filename == 'word/document.xml')

    original = content
    # 替换 thickThinLargeGap → single
    content = content.replace('w:val="thickThinLargeGap"', 'w:val="single"')
    # 替换边框灰色 → auto
    content = re.sub(
        r'(<w:(?:top|bottom|start|end|insideH|insideV)\b[^/]*?)w:color="C0C0C0"',
        r'\1w:color="auto"',
        content
    )

    if content != original:
        count = original.count('thickThinLargeGap')
        with zipfile.ZipFile(docx_path, 'w') as zout:
            zout.writestr(doc_info, content.encode('utf-8'))
            for info, data in entries:
                zout.writestr(info, data)
        print(f'  边框修正：已将 {count} 处 thickThinLargeGap 替换为 single', file=sys.stderr)


def main():
    args = parse_args()

    # 早期验证
    if not args.template.exists():
        print(f'错误：模板文件不存在：{args.template}', file=sys.stderr)
        sys.exit(1)
    if not args.output.parent.exists():
        print(f'错误：输出目录不存在：{args.output.parent}', file=sys.stderr)
        sys.exit(1)

    import docx

    doc = docx.Document(str(args.template))

    field_values = {
        '软件全称': args.name,
        '软件简称': args.short_name,
        '开发完成日期': args.dev_complete_date,
        '首次发表日期': args.first_publish_date,
        '主要功能': args.features,
        '技术特点': args.tech,
    }

    # 上下文关键词 → 字段值映射（与模板中 {变量} 出现顺序一致）
    context_map = [
        ('软件名称', '软件全称', field_values['软件全称']),
        ('软件简称', '软件简称', field_values['软件简称']),
        ('开发完成日期', '开发完成日期', field_values['开发完成日期']),
        ('首次发表日期', '首次发表日期', field_values['首次发表日期']),
        ('主要功能', '主要功能', field_values['主要功能']),
        ('技术特点', '技术特点', field_values['技术特点']),
    ]

    # --- 递归遍历嵌套表格 ---
    def process_table_recursive(table):
        """递归遍历表格，包括嵌套在单元格内的子表格。"""
        count = 0
        for row in table.rows:
            for cell in row.cells:
                count += process_paragraphs(cell.paragraphs, context_map)
                for nested_table in cell.tables:
                    count += process_table_recursive(nested_table)
        return count

    total_replaced = process_paragraphs(doc.paragraphs, context_map)
    for table in doc.tables:
        total_replaced += process_table_recursive(table)

    if total_replaced == EXPECTED_PLACEHOLDERS:
        # python-docx 层面替换成功
        print(
            f'\n共替换 {total_replaced} 个 {{变量}} 占位符（预期为 {EXPECTED_PLACEHOLDERS} 个）',
            file=sys.stderr,
        )
        doc.save(str(args.output))
        print(str(args.output))
    else:
        # python-docx 未完全替换（0 个或部分），回退到 XML 层面替换
        if total_replaced > 0:
            print(
                f'\npython-docx 仅替换 {total_replaced}/{EXPECTED_PLACEHOLDERS} 个，'
                '放弃部分结果，回退到 XML 层面替换...',
                file=sys.stderr,
            )
        else:
            print('\npython-docx 未找到占位符，回退到 XML 层面替换...', file=sys.stderr)
        xml_replaced = xml_level_replace(
            str(args.template), field_values, str(args.output)
        )
        print(
            f'\n共替换 {xml_replaced} 个 {{变量}} 占位符（预期为 {EXPECTED_PLACEHOLDERS} 个）',
            file=sys.stderr,
        )
        if xml_replaced != EXPECTED_PLACEHOLDERS:
            print('替换数量与预期不符，请打开成品文件检查', file=sys.stderr)

    # 修正 LibreOffice 转换导致的边框样式错误
    fix_libreoffice_borders(str(args.output))
    print(str(args.output))


if __name__ == '__main__':
    main()
