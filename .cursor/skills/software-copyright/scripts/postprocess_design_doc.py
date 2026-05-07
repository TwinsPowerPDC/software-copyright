#!/usr/bin/env python3
"""设计文档后处理脚本 — 对 pandoc 生成的 docx 统一字体并修复表格边框。"""

import argparse
import sys
from pathlib import Path


def parse_args():
    parser = argparse.ArgumentParser(
        description='对 pandoc 生成的设计文档 docx 进行后处理：统一字体、修复表格边框'
    )
    parser.add_argument(
        '--input', required=True, type=Path, metavar='DOCX',
        help='输入 .docx 文件路径'
    )
    parser.add_argument(
        '--output', type=Path, metavar='DOCX',
        help='输出 .docx 文件路径（默认覆盖输入文件）'
    )
    return parser.parse_args()


def set_run_font(run, font_name='Microsoft YaHei Light', font_size_pt=None, color_hex='000000'):
    """设置 run 的字体（Latin + 东亚 + CS）和颜色，可选字号。"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run.font.name = font_name
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    # 清除 themeColor（Word 中 themeColor 会覆盖 rgb 值）
    rPr = run._element.get_or_add_rPr()
    color_elem = rPr.find(qn('w:color'))
    if color_elem is not None:
        for attr in ('w:themeColor', 'w:themeTint', 'w:themeShade'):
            if qn(attr) in color_elem.attrib:
                del color_elem.attrib[qn(attr)]
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_style_font(style, font_name='Microsoft YaHei Light', color_hex='000000'):
    """修改样式的字体设置（Latin + 东亚 + CS）和颜色，并清除主题色。"""
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    style.font.name = font_name
    style.font.color.rgb = RGBColor.from_string(color_hex)
    # 清除 themeColor（Word 中 themeColor 会覆盖 rgb 值）
    rPr = style.element.get_or_add_rPr()
    color_elem = rPr.find(qn('w:color'))
    if color_elem is not None:
        for attr in ('w:themeColor', 'w:themeTint', 'w:themeShade'):
            if qn(attr) in color_elem.attrib:
                del color_elem.attrib[qn(attr)]
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def fix_table_borders(table):
    """为表格添加黑色细线边框（0.5pt single，w:sz=4）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # 移除旧边框设置
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)

    # 创建新边框
    tblBorders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')       # 0.5pt
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        tblBorders.append(element)
    tblPr.append(tblBorders)


def process_paragraphs(paragraphs):
    """遍历段落中的所有 run，统一字体。"""
    for para in paragraphs:
        for run in para.runs:
            set_run_font(run)


def process_table_recursive(table):
    """递归处理表格：修复边框 + 统一单元格内字体。"""
    fix_table_borders(table)
    for row in table.rows:
        for cell in row.cells:
            process_paragraphs(cell.paragraphs)
            for nested_table in cell.tables:
                process_table_recursive(nested_table)


def set_doc_defaults(doc, font_name='Microsoft YaHei Light', color_hex='000000'):
    """设置文档级默认字体（w:docDefaults > w:rPrDefault），作为最底层 fallback。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    styles_elem = doc.styles.element
    doc_defaults = styles_elem.find(qn('w:docDefaults'))
    if doc_defaults is None:
        doc_defaults = OxmlElement('w:docDefaults')
        styles_elem.insert(0, doc_defaults)

    rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
    if rPrDefault is None:
        rPrDefault = OxmlElement('w:rPrDefault')
        doc_defaults.insert(0, rPrDefault)

    rPr = rPrDefault.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        rPrDefault.append(rPr)

    # 设置默认字体
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)

    # 设置默认颜色（黑色）并清除主题色
    color_elem = rPr.find(qn('w:color'))
    if color_elem is None:
        color_elem = OxmlElement('w:color')
        rPr.append(color_elem)
    color_elem.set(qn('w:val'), color_hex)
    for attr in ('w:themeColor', 'w:themeTint', 'w:themeShade'):
        if qn(attr) in color_elem.attrib:
            del color_elem.attrib[qn(attr)]


def main():
    args = parse_args()

    if not args.input.exists():
        print(f'错误：输入文件不存在：{args.input}', file=sys.stderr)
        sys.exit(1)

    output_path = args.output or args.input

    import docx

    doc = docx.Document(str(args.input))

    # 0. 设置文档级默认字体（最底层 fallback）
    set_doc_defaults(doc)

    # 1. 遍历所有样式，统一字体和颜色
    for style in doc.styles:
        if hasattr(style, 'font'):
            set_style_font(style)

    # 2. 遍历所有正文段落，逐 run 设置字体
    process_paragraphs(doc.paragraphs)

    # 3. 遍历所有表格：修复边框 + 统一字体
    for table in doc.tables:
        process_table_recursive(table)

    doc.save(str(output_path))
    print(f'后处理完成：{output_path}', file=sys.stderr)
    print(str(output_path))


if __name__ == '__main__':
    main()
